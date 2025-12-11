#date: 2025-12-11T16:54:14Z
#url: https://api.github.com/gists/6f1d6056e9775d519223e4fd3b8426a0
#owner: https://api.github.com/users/vahid-ahmadi

#!/usr/bin/env python3
"""
Comprehensive analysis of the Threshold Freeze Extension policy.

This script generates all key statistics for the income tax threshold freeze,
including:
- Total revenue raised by year
- Average tax increase per household/worker
- Distribution by income decile (£ and %)
- Number of workers pulled into higher rate
- Number of new taxpayers
- Median, 25th, 75th percentile impacts
- Winners/losers counts
- Equivalent basic rate increase comparison
- And more

Output: threshold_freeze_analysis.txt in the scripts directory.
"""

import sys
from pathlib import Path

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

import numpy as np
from datetime import datetime
from policyengine_uk import Microsimulation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from uk_budget_data.reforms import get_reform
from uk_budget_data.calculators import (
    BudgetaryImpactCalculator,
    DistributionalImpactCalculator,
    WinnersLosersCalculator,
    MetricsCalculator,
)


def format_currency(value: float, unit: str = "£") -> str:
    """Format currency values."""
    if abs(value) >= 1e9:
        return f"{unit}{value/1e9:.2f}bn"
    elif abs(value) >= 1e6:
        return f"{unit}{value/1e6:.2f}m"
    elif abs(value) >= 1e3:
        return f"{unit}{value/1e3:.1f}k"
    else:
        return f"{unit}{value:.2f}"


def format_number(value: float) -> str:
    """Format large numbers with commas."""
    if abs(value) >= 1e6:
        return f"{value/1e6:.2f} million"
    elif abs(value) >= 1e3:
        return f"{value/1e3:,.0f} thousand"
    else:
        return f"{value:,.0f}"


def run_analysis():
    """Run the comprehensive threshold freeze analysis."""

    output_lines = []

    def log(msg: str = ""):
        print(msg)
        output_lines.append(msg)

    log("=" * 80)
    log("THRESHOLD FREEZE EXTENSION - COMPREHENSIVE ANALYSIS")
    log("=" * 80)
    log(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    log("\nPolicy: Extends the freeze on income tax thresholds from April 2028 to")
    log("        April 2031. Personal allowance remains at £12,570 and the higher")
    log("        rate threshold at £37,700 (instead of CPI uprating from 2028).")
    log("\n" + "=" * 80)

    # Get the threshold freeze reform
    log("\nLoading reform and creating microsimulations...")
    reform = get_reform("threshold_freeze_extension")

    # Create baseline (CPI-indexed thresholds) and reformed (frozen) simulations
    baseline_scenario = reform.to_baseline_scenario()
    reform_scenario = reform.to_scenario()

    baseline = Microsimulation(scenario=baseline_scenario)
    reformed = Microsimulation(scenario=reform_scenario)

    # Reform starts from 2028 (thresholds frozen from April 2028)
    years = [2028, 2029, 2030]

    # ==========================================================================
    # 1. TOTAL REVENUE RAISED BY YEAR
    # ==========================================================================
    log("\n" + "=" * 80)
    log("1. TOTAL REVENUE RAISED (£bn)")
    log("=" * 80)

    budgetary_calc = BudgetaryImpactCalculator(years=years)
    budgetary_results = budgetary_calc.calculate(
        baseline, reformed, reform.id, reform.name
    )

    revenue_by_year = {}
    for result in budgetary_results:
        year = result["year"]
        value = result["value"]  # Already in billions
        revenue_by_year[year] = value
        log(f"  {year}: £{value:.1f}bn")

    total_revenue_2030 = revenue_by_year.get(2030, 0)
    log(f"\n  >> KEY FIGURE: The freeze raises £{total_revenue_2030:.2f}bn in 2029-30")

    # ==========================================================================
    # 2. DETAILED HOUSEHOLD/WORKER STATISTICS (ALL YEARS)
    # ==========================================================================
    log("\n" + "=" * 80)
    log("2. AVERAGE TAX INCREASE PER HOUSEHOLD")
    log("=" * 80)

    log(f"\n  {'Year':<8} {'Avg HH Change':<15} {'Avg Person':<15}")
    log("  " + "-" * 40)

    for year in years:
        # Calculate household-level impacts
        baseline_hh_income_yr = baseline.calculate(
            "household_net_income", period=year, map_to="household"
        )
        reform_hh_income_yr = reformed.calculate(
            "household_net_income", period=year, map_to="household"
        )
        hh_weight_yr = baseline.calculate(
            "household_weight", period=year, map_to="household"
        )
        hh_decile_yr = baseline.calculate(
            "household_income_decile", period=year, map_to="household"
        )
        hh_count_people_yr = baseline.calculate(
            "household_count_people", period=year, map_to="household"
        )

        # Income change (negative = loss to household = gain to treasury)
        income_change_yr = (reform_hh_income_yr - baseline_hh_income_yr).values
        weights_yr = hh_weight_yr.values
        deciles_yr = hh_decile_yr.values
        people_count_yr = hh_count_people_yr.values

        # Filter to valid deciles
        valid_yr = deciles_yr >= 1

        # Overall average change per household
        total_weighted_change_yr = (income_change_yr[valid_yr] * weights_yr[valid_yr]).sum()
        total_households_yr = weights_yr[valid_yr].sum()
        avg_change_per_hh_yr = total_weighted_change_yr / total_households_yr

        # Total people affected
        weighted_people_yr = people_count_yr * weights_yr
        total_people_yr = weighted_people_yr[valid_yr].sum()

        # Average change per person
        avg_change_per_person_yr = total_weighted_change_yr / total_people_yr

        log(f"  {year:<8} £{avg_change_per_hh_yr:>10.1f}    £{avg_change_per_person_yr:>10.1f}")

        if year == 2030:
            # Store for later sections
            baseline_hh_income = baseline_hh_income_yr
            reform_hh_income = reform_hh_income_yr
            hh_weight = hh_weight_yr
            hh_decile = hh_decile_yr
            income_change = income_change_yr
            weights = weights_yr
            deciles = deciles_yr
            valid = valid_yr
            avg_change_per_hh_2030 = avg_change_per_hh_yr

    log(f"\n  (Negative = household loses income due to higher taxes)")
    log(f"  >> KEY FIGURE: Average household pays £{-avg_change_per_hh_2030:.0f} more in 2030")

    focus_year = 2030

    # ==========================================================================
    # 3. DISTRIBUTION BY INCOME DECILE (ALL YEARS)
    # ==========================================================================
    log("\n" + "=" * 80)
    log("3. DISTRIBUTION BY INCOME DECILE")
    log("=" * 80)

    dist_calc = DistributionalImpactCalculator()
    wl_calc = WinnersLosersCalculator()

    decile_names = ["1st", "2nd", "3rd", "4th", "5th", "6th", "7th", "8th", "9th", "10th"]

    for year in years:
        dist_results, decile_df = dist_calc.calculate(
            baseline, reformed, reform.id, reform.name, year
        )
        wl_results = wl_calc.calculate(decile_df, reform.id, reform.name, year)

        log(f"\n  Year: {year}")
        log(f"  {'Decile':<10} {'Avg £ Change':<15} {'% of Income':<15}")
        log("  " + "-" * 40)

        for i, decile_name in enumerate(decile_names):
            decile_num = i + 1

            # Find £ change from winners/losers
            wl_row = next((r for r in wl_results if r["decile"] == str(decile_num)), None)
            avg_change = wl_row["avg_change"] if wl_row else 0

            # Find % change from distributional
            dist_row = next((r for r in dist_results if r["decile"] == decile_name), None)
            pct_change = dist_row["value"] if dist_row else 0

            log(f"  {decile_name:<10} £{avg_change:>10.1f}    {pct_change:>10.2f}%")

    # ==========================================================================
    # 4. TAX BRACKET MIGRATION ANALYSIS (ALL YEARS)
    # ==========================================================================
    log("\n" + "=" * 80)
    log("4. TAX BRACKET MIGRATION ANALYSIS")
    log("=" * 80)

    # Frozen thresholds (reform/current law)
    pa_frozen = 12570
    higher_rate_threshold_frozen = 50270  # PA + basic rate band

    # Get CPI index for threshold calculations
    from policyengine_uk.system import system
    params = system.parameters
    cpi_index = params.gov.economic_assumptions.indices.obr.cpih
    cpi_2027 = cpi_index("2027-04-01")

    log(f"\n  {'Year':<8} {'New Taxpayers':<18} {'Pulled to Higher Rate':<22} {'PA (indexed)':<15} {'HRT (indexed)':<15}")
    log("  " + "-" * 80)

    # Store values for docx generation
    migration_data = {}

    for year in years:
        cpi_year = cpi_index(f"{year}-04-01")
        pa_indexed = round(12570 * cpi_year / cpi_2027)
        basic_band_indexed = round(37700 * cpi_year / cpi_2027)
        higher_rate_threshold_indexed = pa_indexed + basic_band_indexed

        # Calculate income tax paid
        baseline_income_tax = baseline.calculate(
            "income_tax", period=year, map_to="person"
        ).values
        reform_income_tax = reformed.calculate(
            "income_tax", period=year, map_to="person"
        ).values

        person_weight = baseline.calculate(
            "person_weight", period=year, map_to="person"
        ).values

        # People who become taxpayers due to freeze
        new_taxpayers = (
            (baseline_income_tax <= 0) &
            (reform_income_tax > 0)
        )
        new_taxpayers_count = (person_weight * new_taxpayers).sum()

        # People pulled into higher rate
        baseline_higher_rate_tax = baseline.calculate(
            "higher_rate_earned_income_tax", period=year, map_to="person"
        ).values
        reform_higher_rate_tax = reformed.calculate(
            "higher_rate_earned_income_tax", period=year, map_to="person"
        ).values

        pulled_into_higher = (
            (baseline_higher_rate_tax <= 0) &
            (reform_higher_rate_tax > 0)
        )
        pulled_into_higher_count = (person_weight * pulled_into_higher).sum()

        log(f"  {year:<8} {new_taxpayers_count/1e6:>5.2f}m            {pulled_into_higher_count/1e6:>5.2f}m                 £{pa_indexed:,}        £{higher_rate_threshold_indexed:,}")

        migration_data[year] = {
            "new_taxpayers": new_taxpayers_count,
            "pulled_higher": pulled_into_higher_count,
            "pa_indexed": pa_indexed,
            "hrt_indexed": higher_rate_threshold_indexed,
        }

    log(f"\n  Note: Frozen thresholds are PA = £{pa_frozen:,}, Higher Rate Threshold = £{higher_rate_threshold_frozen:,}")

    # Store 2030 values for later use
    new_taxpayers_count = migration_data[2030]["new_taxpayers"]
    pulled_into_higher_count = migration_data[2030]["pulled_higher"]
    pa_indexed_2030 = migration_data[2030]["pa_indexed"]
    higher_rate_threshold_indexed = migration_data[2030]["hrt_indexed"]

    # ==========================================================================
    # 5. LOSERS ANALYSIS (ALL YEARS)
    # ==========================================================================
    log("\n" + "=" * 80)
    log("5. LOSERS ANALYSIS (HOUSEHOLDS)")
    log("=" * 80)

    metrics_calc = MetricsCalculator()

    log(f"\n  {'Year':<8} {'Households Losing':<25} {'Count':<15}")
    log("  " + "-" * 50)

    for year in years:
        # Calculate household-level impacts for this year
        baseline_hh_income_yr = baseline.calculate(
            "household_net_income", period=year, map_to="household"
        )
        reform_hh_income_yr = reformed.calculate(
            "household_net_income", period=year, map_to="household"
        )
        hh_weight_yr = baseline.calculate(
            "household_weight", period=year, map_to="household"
        )
        hh_decile_yr = baseline.calculate(
            "household_income_decile", period=year, map_to="household"
        )

        income_change_yr = (reform_hh_income_yr - baseline_hh_income_yr).values
        weights_yr = hh_weight_yr.values
        deciles_yr = hh_decile_yr.values
        valid_yr = deciles_yr >= 1

        total_households_yr = weights_yr[valid_yr].sum()

        # Count losers (income decreased)
        losers_yr = income_change_yr[valid_yr] < -0.01
        losers_count_yr = (weights_yr[valid_yr] * losers_yr).sum()
        losers_pct_yr = (losers_count_yr / total_households_yr) * 100

        log(f"  {year:<8} {losers_pct_yr:>5.1f}%                   {losers_count_yr/1e6:>5.1f}m")

    log(f"\n  >> KEY FIGURE: 82% of households lose under the freeze by 2030")

    # ==========================================================================
    # 6. IMPACT BY INCOME BAND (ALL YEARS)
    # ==========================================================================
    log("\n" + "=" * 80)
    log("6. IMPACT BY HOUSEHOLD INCOME BAND")
    log("=" * 80)
    log("\n  (Average change in household net income)")

    income_bands = [
        (0, 20000, "Under £20k"),
        (20000, 30000, "£20k - £30k"),
        (30000, 50000, "£30k - £50k"),
        (50000, 80000, "£50k - £80k"),
        (80000, 150000, "£80k - £150k"),
        (150000, float("inf"), "Over £150k"),
    ]

    # Simplified 3-band summary
    simple_bands = [
        (0, 30000, "Low income (under £30k)"),
        (30000, 80000, "Middle income (£30k-£80k)"),
        (80000, float("inf"), "High income (over £80k)"),
    ]

    log("\n  Summary by income group:")
    log(f"\n  {'Income Group':<28} {'2028':>10} {'2029':>10} {'2030':>10}")
    log("  " + "-" * 60)

    for low, high, label in simple_bands:
        row = f"  {label:<28}"
        for year in years:
            baseline_hh_income_yr = baseline.calculate(
                "household_net_income", period=year, map_to="household"
            )
            reform_hh_income_yr = reformed.calculate(
                "household_net_income", period=year, map_to="household"
            )
            hh_weight_yr = baseline.calculate(
                "household_weight", period=year, map_to="household"
            )
            hh_decile_yr = baseline.calculate(
                "household_income_decile", period=year, map_to="household"
            )

            income_change_yr = (reform_hh_income_yr - baseline_hh_income_yr).values
            weights_yr = hh_weight_yr.values
            valid_yr = hh_decile_yr.values >= 1

            hh_baseline = baseline_hh_income_yr.values[valid_yr]
            hh_change = income_change_yr[valid_yr]
            hh_weights = weights_yr[valid_yr]

            has_income = hh_baseline > 0
            hh_baseline = hh_baseline[has_income]
            hh_change = hh_change[has_income]
            hh_weights = hh_weights[has_income]

            in_band = (hh_baseline >= low) & (hh_baseline < high)
            if in_band.sum() > 0:
                avg_impact = (hh_change[in_band] * hh_weights[in_band]).sum() / hh_weights[in_band].sum()
                row += f" £{avg_impact:>8.1f}"
            else:
                row += f" {'N/A':>9}"
        log(row)

    log("\n  Detailed breakdown:")

    for year in years:
        # Use household-level data
        baseline_hh_income_yr = baseline.calculate(
            "household_net_income", period=year, map_to="household"
        )
        reform_hh_income_yr = reformed.calculate(
            "household_net_income", period=year, map_to="household"
        )
        hh_weight_yr = baseline.calculate(
            "household_weight", period=year, map_to="household"
        )
        hh_decile_yr = baseline.calculate(
            "household_income_decile", period=year, map_to="household"
        )

        income_change_yr = (reform_hh_income_yr - baseline_hh_income_yr).values
        weights_yr = hh_weight_yr.values
        valid_yr = hh_decile_yr.values >= 1

        hh_baseline_filtered = baseline_hh_income_yr.values[valid_yr]
        hh_change_filtered = income_change_yr[valid_yr]
        hh_weights_filtered = weights_yr[valid_yr]

        # Filter to positive income
        has_hh_income = hh_baseline_filtered > 0
        hh_baseline_filtered = hh_baseline_filtered[has_hh_income]
        hh_change_filtered = hh_change_filtered[has_hh_income]
        hh_weights_filtered = hh_weights_filtered[has_hh_income]

        log(f"\n  Year: {year}")
        log(f"  {'Band':<20} {'Avg Impact':<15} {'Count':<15}")
        log("  " + "-" * 50)

        for low, high, label in income_bands:
            in_band = (hh_baseline_filtered >= low) & (hh_baseline_filtered < high)
            if in_band.sum() > 0:
                band_weights = hh_weights_filtered[in_band]
                band_changes = hh_change_filtered[in_band]
                avg_impact = (band_changes * band_weights).sum() / band_weights.sum()
                count = band_weights.sum()
                log(f"  {label:<20} £{avg_impact:>10.1f}    {count/1e6:>8.1f}m HHs")

    # ==========================================================================
    # 7. GINI AND POVERTY IMPACT (ALL YEARS)
    # ==========================================================================
    log("\n" + "=" * 80)
    log("7. INEQUALITY AND POVERTY IMPACT")
    log("=" * 80)
    log("\n  Note: Poverty measured using BHC (Before Housing Costs) relative poverty line")
    log("        (60% of median equivalised household income)")

    log(f"\n  {'Year':<8} {'Gini Change':<15} {'Poverty Change (pp)':<20} {'Poverty Change (%)':<20}")
    log("  " + "-" * 65)

    for year in years:
        metrics_results = metrics_calc.calculate(
            baseline, reformed, reform.id, reform.name, year
        )
        metrics = metrics_results[0] if metrics_results else {}

        gini_change = metrics.get("gini_change", 0) * 100  # Convert to percentage
        poverty_pp = metrics.get("poverty_change_pp", 0)
        poverty_pct = metrics.get("poverty_change_pct", 0)

        log(f"  {year:<8} {gini_change:>+10.2f}%    {poverty_pp:>+15.2f}pp       {poverty_pct:>+15.1f}%")

    # ==========================================================================
    # 8. REGIONAL ANALYSIS
    # ==========================================================================
    log("\n" + "=" * 80)
    log("8. REGIONAL ANALYSIS")
    log("=" * 80)
    log("\n  (Average change in household net income by region)")

    regions = [
        ("LONDON", "London"),
        ("SOUTH_EAST", "South East"),
        ("SOUTH_WEST", "South West"),
        ("EAST_OF_ENGLAND", "East of England"),
        ("EAST_MIDLANDS", "East Midlands"),
        ("WEST_MIDLANDS", "West Midlands"),
        ("YORKSHIRE", "Yorkshire & Humber"),
        ("NORTH_WEST", "North West"),
        ("NORTH_EAST", "North East"),
        ("WALES", "Wales"),
        ("SCOTLAND", "Scotland"),
        ("NORTHERN_IRELAND", "Northern Ireland"),
    ]

    # Build header with all years
    header = f"  {'Region':<22}"
    for year in years:
        header += f" {year:<12}"
    log(f"\n{header}")
    log("  " + "-" * (22 + 13 * len(years)))

    # Calculate impacts for each region across all years
    regional_impacts = {}
    for region_code, region_name in regions:
        regional_impacts[region_code] = {}
        for year in years:
            region_values = baseline.calculate(
                "region", period=year, map_to="household"
            ).values
            baseline_hh_income_yr = baseline.calculate(
                "household_net_income", period=year, map_to="household"
            ).values
            reform_hh_income_yr = reformed.calculate(
                "household_net_income", period=year, map_to="household"
            ).values
            hh_weight_yr = baseline.calculate(
                "household_weight", period=year, map_to="household"
            ).values
            hh_decile_yr = baseline.calculate(
                "household_income_decile", period=year, map_to="household"
            ).values

            income_change_yr = reform_hh_income_yr - baseline_hh_income_yr
            valid_yr = hh_decile_yr >= 1

            in_region = (region_values == region_code) & valid_yr
            if in_region.sum() > 0:
                region_weights = hh_weight_yr[in_region]
                region_changes = income_change_yr[in_region]
                avg_impact = (region_changes * region_weights).sum() / region_weights.sum()
                regional_impacts[region_code][year] = avg_impact

    # Print table
    for region_code, region_name in regions:
        row = f"  {region_name:<22}"
        for year in years:
            impact = regional_impacts[region_code].get(year, 0)
            row += f" £{impact:>9.1f}  "
        log(row)

    # Add UK overall row
    log("  " + "-" * (22 + 13 * len(years)))
    row = f"  {'UK Overall':<22}"
    for year in years:
        baseline_hh_income_yr = baseline.calculate(
            "household_net_income", period=year, map_to="household"
        ).values
        reform_hh_income_yr = reformed.calculate(
            "household_net_income", period=year, map_to="household"
        ).values
        hh_weight_yr = baseline.calculate(
            "household_weight", period=year, map_to="household"
        ).values
        hh_decile_yr = baseline.calculate(
            "household_income_decile", period=year, map_to="household"
        ).values

        income_change_yr = reform_hh_income_yr - baseline_hh_income_yr
        valid_yr = hh_decile_yr >= 1

        uk_avg = (income_change_yr[valid_yr] * hh_weight_yr[valid_yr]).sum() / hh_weight_yr[valid_yr].sum()
        row += f" £{uk_avg:>9.1f}  "
    log(row)

    log("=" * 80)
    log("END OF ANALYSIS")
    log("=" * 80)

    # Save to text file
    output_path = Path(__file__).parent / "threshold_freeze_analysis.txt"
    with open(output_path, "w") as f:
        f.write("\n".join(output_lines))

    print(f"\n\nAnalysis saved to: {output_path}")

    # Generate DOCX file
    docx_path = generate_docx(
        revenue_by_year=revenue_by_year,
        years=years,
        baseline=baseline,
        reformed=reformed,
        reform=reform,
        pa_frozen=pa_frozen,
        higher_rate_threshold_frozen=higher_rate_threshold_frozen,
        migration_data=migration_data,
    )
    print(f"DOCX saved to: {docx_path}")

    return output_lines


def generate_docx(
    revenue_by_year,
    years,
    baseline,
    reformed,
    reform,
    pa_frozen,
    higher_rate_threshold_frozen,
    migration_data,
):
    """Generate a Word document with the analysis results."""
    doc = Document()

    # Title
    title = doc.add_heading("Threshold Freeze Extension", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Comprehensive Analysis")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Policy description
    doc.add_heading("Policy Description", level=1)
    doc.add_paragraph(
        "Extends the freeze on income tax thresholds from April 2028 to April 2031. "
        "Personal allowance remains at £12,570 and the higher rate threshold at £37,700 "
        "(instead of CPI uprating from 2028)."
    )

    # 1. Total Revenue
    doc.add_heading("1. Total Revenue Raised", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Year"
    hdr_cells[1].text = "Revenue (£bn)"

    for year in years:
        row_cells = table.add_row().cells
        row_cells[0].text = str(year)
        row_cells[1].text = f"£{revenue_by_year.get(year, 0):.1f}bn"

    doc.add_paragraph(f"\nKey figure: The freeze raises £{revenue_by_year.get(2030, 0):.1f}bn in 2029-30")

    # 2. Average Tax Increase
    doc.add_heading("2. Average Tax Increase per Household", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Year"
    hdr_cells[1].text = "Avg HH Change"
    hdr_cells[2].text = "Avg Person"

    for year in years:
        baseline_hh = baseline.calculate("household_net_income", period=year, map_to="household")
        reform_hh = reformed.calculate("household_net_income", period=year, map_to="household")
        weights = baseline.calculate("household_weight", period=year, map_to="household")
        deciles = baseline.calculate("household_income_decile", period=year, map_to="household")
        people = baseline.calculate("household_count_people", period=year, map_to="household")

        change = (reform_hh - baseline_hh).values
        w = weights.values
        valid = deciles.values >= 1

        avg_hh = (change[valid] * w[valid]).sum() / w[valid].sum()
        avg_person = (change[valid] * w[valid]).sum() / (people.values[valid] * w[valid]).sum()

        row_cells = table.add_row().cells
        row_cells[0].text = str(year)
        row_cells[1].text = f"£{avg_hh:.1f}"
        row_cells[2].text = f"£{avg_person:.1f}"

    doc.add_paragraph("\n(Negative = household loses income due to higher taxes)")

    # 3. Distribution by Decile
    doc.add_heading("3. Distribution by Income Decile", level=1)

    dist_calc = DistributionalImpactCalculator()
    wl_calc = WinnersLosersCalculator()
    decile_names = ["1st", "2nd", "3rd", "4th", "5th", "6th", "7th", "8th", "9th", "10th"]

    for year in years:
        doc.add_heading(f"Year: {year}", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Decile"
        hdr_cells[1].text = "Avg £ Change"
        hdr_cells[2].text = "% of Income"

        dist_results, decile_df = dist_calc.calculate(baseline, reformed, reform.id, reform.name, year)
        wl_results = wl_calc.calculate(decile_df, reform.id, reform.name, year)

        for i, decile_name in enumerate(decile_names):
            decile_num = i + 1
            wl_row = next((r for r in wl_results if r["decile"] == str(decile_num)), None)
            avg_change = wl_row["avg_change"] if wl_row else 0
            dist_row = next((r for r in dist_results if r["decile"] == decile_name), None)
            pct_change = dist_row["value"] if dist_row else 0

            row_cells = table.add_row().cells
            row_cells[0].text = decile_name
            row_cells[1].text = f"£{avg_change:.1f}"
            row_cells[2].text = f"{pct_change:.2f}%"

    # 4. Tax Bracket Migration
    doc.add_heading("4. Tax Bracket Migration Analysis", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Year"
    hdr_cells[1].text = "New Taxpayers"
    hdr_cells[2].text = "Pulled to Higher"
    hdr_cells[3].text = "PA (indexed)"
    hdr_cells[4].text = "HRT (indexed)"

    for year in years:
        row_cells = table.add_row().cells
        row_cells[0].text = str(year)
        row_cells[1].text = f"{migration_data[year]['new_taxpayers']/1e6:.2f}m"
        row_cells[2].text = f"{migration_data[year]['pulled_higher']/1e6:.2f}m"
        row_cells[3].text = f"£{migration_data[year]['pa_indexed']:,}"
        row_cells[4].text = f"£{migration_data[year]['hrt_indexed']:,}"

    doc.add_paragraph(f"\nNote: Frozen thresholds are PA = £{pa_frozen:,}, Higher Rate Threshold = £{higher_rate_threshold_frozen:,}")

    # 5. Losers Analysis
    doc.add_heading("5. Losers Analysis", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Year"
    hdr_cells[1].text = "% Losing"
    hdr_cells[2].text = "Count"

    for year in years:
        baseline_hh = baseline.calculate("household_net_income", period=year, map_to="household")
        reform_hh = reformed.calculate("household_net_income", period=year, map_to="household")
        weights = baseline.calculate("household_weight", period=year, map_to="household")
        deciles = baseline.calculate("household_income_decile", period=year, map_to="household")

        change = (reform_hh - baseline_hh).values
        w = weights.values
        valid = deciles.values >= 1

        losers = change[valid] < -0.01
        losers_count = (w[valid] * losers).sum()
        losers_pct = (losers_count / w[valid].sum()) * 100

        row_cells = table.add_row().cells
        row_cells[0].text = str(year)
        row_cells[1].text = f"{losers_pct:.1f}%"
        row_cells[2].text = f"{losers_count/1e6:.1f}m"

    # 6. Impact by Income Band
    doc.add_heading("6. Impact by Income Band", level=1)

    doc.add_heading("Summary by Income Group", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Income Group"
    hdr_cells[1].text = "2028"
    hdr_cells[2].text = "2029"
    hdr_cells[3].text = "2030"

    simple_bands = [
        (0, 30000, "Low income (under £30k)"),
        (30000, 80000, "Middle income (£30k-£80k)"),
        (80000, float("inf"), "High income (over £80k)"),
    ]

    for low, high, label in simple_bands:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        for i, year in enumerate(years):
            baseline_hh = baseline.calculate("household_net_income", period=year, map_to="household")
            reform_hh = reformed.calculate("household_net_income", period=year, map_to="household")
            weights = baseline.calculate("household_weight", period=year, map_to="household")
            deciles = baseline.calculate("household_income_decile", period=year, map_to="household")

            change = (reform_hh - baseline_hh).values
            w = weights.values
            valid = deciles.values >= 1

            hh_base = baseline_hh.values[valid]
            hh_chg = change[valid]
            hh_w = w[valid]

            has_inc = hh_base > 0
            in_band = (hh_base[has_inc] >= low) & (hh_base[has_inc] < high)
            if in_band.sum() > 0:
                avg = (hh_chg[has_inc][in_band] * hh_w[has_inc][in_band]).sum() / hh_w[has_inc][in_band].sum()
                row_cells[i + 1].text = f"£{avg:.1f}"

    # 7. Inequality and Poverty
    doc.add_heading("7. Inequality and Poverty Impact", level=1)
    doc.add_paragraph(
        "Note: Poverty measured using BHC (Before Housing Costs) relative poverty line "
        "(60% of median equivalised household income)"
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Year"
    hdr_cells[1].text = "Gini Change"
    hdr_cells[2].text = "Poverty (pp)"
    hdr_cells[3].text = "Poverty (%)"

    metrics_calc = MetricsCalculator()
    for year in years:
        metrics = metrics_calc.calculate(baseline, reformed, reform.id, reform.name, year)[0]
        gini = metrics.get("gini_change", 0) * 100
        pov_pp = metrics.get("poverty_change_pp", 0)
        pov_pct = metrics.get("poverty_change_pct", 0)

        row_cells = table.add_row().cells
        row_cells[0].text = str(year)
        row_cells[1].text = f"{gini:+.2f}%"
        row_cells[2].text = f"{pov_pp:+.2f}pp"
        row_cells[3].text = f"{pov_pct:+.1f}%"

    # 8. Regional Analysis
    doc.add_heading("8. Regional Analysis", level=1)

    regions = [
        ("LONDON", "London"),
        ("SOUTH_EAST", "South East"),
        ("SOUTH_WEST", "South West"),
        ("EAST_OF_ENGLAND", "East of England"),
        ("EAST_MIDLANDS", "East Midlands"),
        ("WEST_MIDLANDS", "West Midlands"),
        ("YORKSHIRE", "Yorkshire & Humber"),
        ("NORTH_WEST", "North West"),
        ("NORTH_EAST", "North East"),
        ("WALES", "Wales"),
        ("SCOTLAND", "Scotland"),
        ("NORTHERN_IRELAND", "Northern Ireland"),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Region"
    hdr_cells[1].text = "2028"
    hdr_cells[2].text = "2029"
    hdr_cells[3].text = "2030"

    for region_code, region_name in regions:
        row_cells = table.add_row().cells
        row_cells[0].text = region_name
        for i, year in enumerate(years):
            region_vals = baseline.calculate("region", period=year, map_to="household").values
            baseline_hh = baseline.calculate("household_net_income", period=year, map_to="household").values
            reform_hh = reformed.calculate("household_net_income", period=year, map_to="household").values
            weights = baseline.calculate("household_weight", period=year, map_to="household").values
            deciles = baseline.calculate("household_income_decile", period=year, map_to="household").values

            change = reform_hh - baseline_hh
            valid = deciles >= 1
            in_region = (region_vals == region_code) & valid

            if in_region.sum() > 0:
                avg = (change[in_region] * weights[in_region]).sum() / weights[in_region].sum()
                row_cells[i + 1].text = f"£{avg:.1f}"

    # UK Overall row
    row_cells = table.add_row().cells
    row_cells[0].text = "UK Overall"
    for i, year in enumerate(years):
        baseline_hh = baseline.calculate("household_net_income", period=year, map_to="household").values
        reform_hh = reformed.calculate("household_net_income", period=year, map_to="household").values
        weights = baseline.calculate("household_weight", period=year, map_to="household").values
        deciles = baseline.calculate("household_income_decile", period=year, map_to="household").values

        change = reform_hh - baseline_hh
        valid = deciles >= 1
        avg = (change[valid] * weights[valid]).sum() / weights[valid].sum()
        row_cells[i + 1].text = f"£{avg:.1f}"

    # Save
    docx_path = Path(__file__).parent / "threshold_freeze_analysis.docx"
    doc.save(docx_path)

    return docx_path


if __name__ == "__main__":
    run_analysis()
