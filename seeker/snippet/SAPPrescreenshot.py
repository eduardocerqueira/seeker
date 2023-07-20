#date: 2023-07-20T17:00:16Z
#url: https://api.github.com/gists/d97cc37c358fb569d3db1f3d88d72ed8
#owner: https://api.github.com/users/daggupati45

from docx import Document
from docx.shared import Inches
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait

# Create a new Word document
doc = Document()

driver = webdriver.Chrome()
# chrome_options = Options()
# chrome_options.add_argument('--no-sandbox')
# chrome_options.add_argument('--disable-gpu')
# chrome_options.add_argument('--disable-dev-shm-usage')
# chrome_options.add_argument('--remote-debugging-port=9222')
# chrome_options.add_argument('--incognito')
# chrome_options.add_argument('--headless')
# chrome_options.add_argument("--window-size=1920,1080")
# chrome_options.binary_location = '/usr/bin/google-chrome'
# driver = webdriver.Chrome(chrome_options=chrome_options)

driver.get(
    "https://vlss1ap00a.scmnp.aws.corning.com:44300/sap/bc/gui/sap/its/webgui")



# Find the login fields and enter username and password
def login():
    username_field = driver.find_element(By.ID, "sap-user")
    #    username_field.send_keys("aimbashir")
    username_field.send_keys("cbasis")
    password_field = "**********"
    #   password_field.send_keys("Corning000")
    password_field.send_keys("Corning@2023")  
    #    password_field.send_keys("Corning@2023")   
    #   password_field.send_keys("MDPCorning@23")  
    driver.find_element(By.ID, "LOGON_BUTTON").click()
    WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.ID, 'ToolbarOkCode')))
    time.sleep(5)
    print("Login Successful")

## Capturing the screenshots into word document
    
    driver.find_element(By.ID, "ToolbarOkCode").send_keys("sm21")
    # 10 | sendKeys | id=ToolbarOkCode | ${KEY_ENTER}
    driver.find_element(By.ID, "ToolbarOkCode").send_keys(Keys.ENTER)
    # 11 | click | id=ToolbarOkCode |
    driver.find_element(By.ID, "ToolbarOkCode").click()
    # 12 | type | id=ToolbarOkCode | /n


login()

image_path = 'Image1.png'
##doc.add_paragraph(a)
driver.save_screenshot(image_path)
doc.add_picture(image_path, width=Inches(6))
doc.save('Document_SS1_prescreen.docx')'Document_SS1_prescreen.docx')