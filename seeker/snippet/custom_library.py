#date: 2024-07-18T16:51:36Z
#url: https://api.github.com/gists/59ee700d963b60e8f97a45283cf00ad5
#owner: https://api.github.com/users/BorisITZaitsev

import docx
import json
import os.path


path = os.getcwd().replace('\\', "/") + "/"


def folder(g_id):
    global path
    os.mkdir(path + str(g_id))


def database_create(g_id):
    global path
    if not os.path.exists(path + str(g_id) + "/profiles"):
        doc, data = docx.Document(path + str(g_id) + "/st_list.docx"), {}
        for paragraph in doc.paragraphs[1:]:
            data["".join(paragraph.text)] = [0]
        with open(path + str(g_id) + '/profiles1.json', 'w', encoding="utf-8") as file:
            json.dump(data, file, ensure_ascii=False)
        with open(path + str(g_id) + '/further_speakers.json', 'w', encoding="utf-8") as file:
            json.dump({}, file, ensure_ascii=False)


def single_apply(change, g_id):
    global path
    change = change.split(" - ")
    speaker, theme = change[0], change[1]
    data = get_full_data(g_id)
    data[speaker][0] += 1
    data[speaker].append(theme)
    with open(path + str(g_id) + "/profiles1.json", "w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False)
    with open(path + str(g_id) + "/further_speakers.json", "r", encoding="utf-8") as file:
        data = json.load(file)
        data[speaker] = theme
    with open(path + str(g_id) + "/further_speakers.json", "w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False)


def connector(g_id):
    global path
    data = get_full_data(g_id)
    doc = docx.Document(path + str(g_id) + "/th_list.docx")
    text = ["".join(paragraph.text) for paragraph in doc.paragraphs[2:]]
    new = applier(data, text)
    with open(path + str(g_id) + "/profiles1.json", "w", encoding="utf-8") as file:
        json.dump(new[0], file, ensure_ascii=False)
    with open(path + str(g_id) + "/further_speakers.json", "w", encoding="utf-8") as f2:
        json.dump(new[1], f2, ensure_ascii=False)


def applier(base, array):
    malo = min([base[i][0] for i in base])
    dublicator = [i for i in base]
    speakers, i, b = {}, 0, 0
    while b != len(array):
        if base[dublicator[i]][0] == malo:
            base[dublicator[i]][0] += 1
            base[dublicator[i]].append(array[b])
            if dublicator[i] in speakers:
                theme = speakers[dublicator[i]]
                speakers[dublicator[i]] = [theme, array[b]]
            else:
                speakers[dublicator[i]] = array[b]
            b += 1
        if i == len(dublicator) - 1:
            i = -1
            malo = min([base[i][0] for i in base])
        i += 1
    return base, speakers


def get_full_data(g_id):
    global path
    with open(path + str(g_id) + "/profiles1.json", "r", encoding="utf-8") as base:
        data = json.load(base)
        return data


def base_existence(g_id):
    global path
    return os.path.exists(path + str(g_id))


def further_speakers(g_id):
    global path
    with open(path + str(g_id) + "/further_speakers.json", "r", encoding="utf-8") as file:
        return json.load(file)


def theme_remove(g_id, student, number):
    data = get_full_data(g_id)
    theme = data[student][number]
    data[student].remove(theme)
    data[student][0] = data[student][0] - 1
    with open(path + str(g_id) + "/profiles1.json", "w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False)
    return theme


def hm_themes(g_id):
    data = get_full_data(g_id)
    number = 0
    for i in data:
        number += data[i][0]
    return number


def clear(g_id):
    global path
    with open(path + str(g_id) + "/further_speakers.json", "w", encoding="utf-8") as file:
        json.dump({}, file, ensure_ascii=False)
