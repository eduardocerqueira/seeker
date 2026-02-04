#date: 2026-02-04T17:37:40Z
#url: https://api.github.com/gists/a659ef075ff205d0fe43fcbb475e2700
#owner: https://api.github.com/users/me-suzy

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Titlu principal
title = doc.add_heading('REZOLVAREA PROBLEMELOR LOMBARE ȘI DISFUNCȚIEI ERECTILE PRIN SĂRITURI VERTICALE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_format = title.runs[0].font
title_format.size = Pt(16)
title_format.bold = True

# Subtitlu
subtitle = doc.add_paragraph('Cum am descoperit că săriturile ușoare pe loc sunt mai benefice decât alergarea pentru sănătatea coloanei și funcția sexuală')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0].font
subtitle_format.size = Pt(12)
subtitle_format.italic = True

doc.add_paragraph()

# INTRODUCERE
doc.add_heading('INTRODUCERE: PROBLEMA IDENTIFICATĂ', 1)

p1 = doc.add_paragraph()
p1.add_run('După ani de alergare intensă la pace rapid (4:00-4:30/km), fără un program consistent de stretching, am dezvoltat o serie de simptome aparent neconectate: ').font.size = Pt(11)
p1.add_run('disconfort lombar persistent, disfuncție erectilă graduală, și un simptom neobișnuit - testiculele care se strângeau involuntar după doar 3 minute de alergare.').font.size = Pt(11)

p2 = doc.add_paragraph()
p2.add_run('Consultările medicale au identificat o hernie sau protruzie discală probabilă la nivel L4-L5 sau L5-S1, rezultatul unei leziuni acute suferite într-un concurs de alergare (coborâre rapidă cu hiperextensie lombară). Dar medicamentele anti-inflamatoare (Roboflex) și yoga-ul ameliorau doar temporar simptomele, fără să rezolve cauza de fond.').font.size = Pt(11)

p3 = doc.add_paragraph()
p3.add_run('Descoperirea accidentală a efectului mașinii - unde aveam erecții puternice ca pasager, dar nu acasă - m-a condus către înțelegerea mecanismului real: ').font.size = Pt(11)
p3.add_run('vibrațiile constante (5-15 Hz) combinate cu poziția complet relaxată a picioarelor generau decompresie lombară continuă și relaxare completă a psoas-ului.').font.size = Pt(11)
p3.runs[-1].font.bold = True

# CAUZA REALĂ
doc.add_heading('CAUZA REALĂ: PSOAS CONTRACTAT CRONIC', 1)

p4 = doc.add_paragraph()
p4.add_run('Ani de alergare intensă fără stretching adecvat au scurtat progresiv mușchiul psoas (iliopsoas) - principalul flexor de șold. Acest mușchi se contractă la fiecare pas pentru a ridica genunchiul, iar la pace rapid, se contractă 180 de ori pe minut. Fără stretching post-alergare, psoas-ul rămâne în stare de contractură cronică, trăgând constant pe vertebrele lombare L1-L2.').font.size = Pt(11)

p5 = doc.add_paragraph()
p5.add_run('Mecanismul cascadei de simptome:').font.size = Pt(11)
p5.runs[0].font.bold = True

lista_cauze = [
    'Psoas contractat cronic trage pe L1-L2 anterior',
    'Compresie progresivă a nervului genitofemoral (L1-L2)',
    'Reflex cremasteric exagerat → testiculele se strâng involuntar',
    'Compresie discului herniat pe nervii sacrali S2-S4',
    'Afectare funcție erectilă (nervii S2-S4 controlează erecția)',
    'Lordoză lombară accentuată → compresie cronică disc'
]

for cauza in lista_cauze:
    p = doc.add_paragraph(cauza, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.runs[0].font.size = Pt(11)

# DESCOPERIREA CHEIE
doc.add_heading('DESCOPERIREA CHEIE: SĂRITURI PE LOC VS. ALERGARE', 1)

p6 = doc.add_paragraph()
p6.add_run('Observația care a schimbat totul: ').font.size = Pt(11)
p6.add_run('când săream pe loc (ambele picioare simultan, 10cm înălțime, pe suprafață moale), testiculele rămâneau complet relaxate. Dar când alergam - chiar și pe plajă, chiar la pace ușor - după 3 minute testiculele se strângeau.').font.size = Pt(11)
p6.runs[-1].font.bold = True

p7 = doc.add_paragraph()
p7.add_run('Diferența nu era impactul (plaja e moale), ci ').font.size = Pt(11)
p7.add_run('mecanismul biomecanic fundamental diferit:').font.size = Pt(11)
p7.runs[-1].font.bold = True

# Tabel comparativ - CORECTAT
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'SĂRITURI PE LOC'
hdr_cells[1].text = 'ALERGARE'

# Adăugăm rânduri unul câte unul
data = [
    ('Vector forță: 100% VERTICAL (sus-jos)', 'Vector forță: 70% ORIZONTAL + 30% vertical'),
    ('Decompresie pură în faza de zbor', 'Shear force (forfecare) pe discuri'),
    ('ZERO hip flexion alternativ', 'Hip flexion repetitiv 180x/minut'),
    ('Psoas PASIV, relaxat', 'Psoas ACTIV, contractat intens'),
    ('Ambele picioare simultan = simetric', 'Un picior după altul = asimetric'),
    ('Testiculele RELAXATE', 'Testiculele se STRÂNG (după 3 min)')
]

for col1, col2 in data:
    row_cells = table.add_row().cells
    row_cells[0].text = col1
    row_cells[1].text = col2

doc.add_paragraph()

# MECANISMUL SĂRITURILOR
doc.add_heading('MECANISMUL TERAPEUTIC AL SĂRITURILOR', 1)

p8 = doc.add_paragraph()
p8.add_run('Discurile intervertebrale conțin nucleus pulposus - un gel format din 80-88% apă plus proteoglicani. Acest "gel magic" funcționează ca un amortizor hidraulic, menținând spațiul între vertebre. Dar în cursul zilei, compresia constantă (din statul în picioare, șezut, alergare) elimină lichidul din disc. La 40 de ani, discurile pierd 10-20% din capacitatea de hidratare.').font.size = Pt(11)

p9 = doc.add_paragraph()
p9.add_run('Săriturile verticale creează un mecanism unic de "pompare" a lichidului:').font.size = Pt(11)
p9.runs[0].font.bold = True

doc.add_paragraph()

mecanisme = [
    ('FAZA 1 - ATERIZARE (compresie):', 'Greutatea corpului comprimă discul → lichidul este împins afară în cartilajul terminal și vasele sanguine → annulus fibrosus (inelele de colagen) se întinde elastic'),
    ('FAZA 2 - ZBOR (decompresie):', 'Gravitația trage corpul jos → vertebrele se depărtează → presiunea în disc scade → se creează un vacuum → lichidul este ABSORBIT ÎNAPOI → discul se "reumflă"'),
    ('REPETARE 600-1800x/zi:', 'Fiecare ciclu pompează lichid → rehidratare progresivă → nucleus pulposus devine mai spongy → vertebrele rămân depărtate → nervii (L1-L2, S2-S4) sunt liberi')
]

for titlu, continut in mecanisme:
    p = doc.add_paragraph()
    p.add_run(titlu).font.bold = True
    p.add_run(' ' + continut)
    p.runs[0].font.size = Pt(11)
    p.runs[1].font.size = Pt(11)

# ELASTICITATE VS FLEXIBILITATE
doc.add_heading('ELASTICITATE VS. FLEXIBILITATE: DIFERENȚA CRITICĂ', 1)

p10 = doc.add_paragraph()
p10.add_run('Un antrenor personal mi-a explicat diferența fundamentală dintre aceste concepte, esențială pentru alergători:').font.size = Pt(11)

doc.add_paragraph()

p11 = doc.add_paragraph()
p11.add_run('FLEXIBILITATEA').font.bold = True
p11.add_run(' = range of motion pasiv al unei articulații sau mușchi. Se măsoară static (ex: cât de jos ajungi când te apleci să atingi degetele). Importantă pentru prevenirea leziunilor, dar INSUFICIENTĂ pentru alergători.').font.size = Pt(11)

p12 = doc.add_paragraph()
p12.add_run('ELASTICITATEA').font.bold = True
p12.add_run(' = capacitatea țesutului de a stoca și elibera energie rapid, ca un arc sau elastic. Tendoanele (Achiles, patellar) și fascia funcționează ca arcuri la fiecare pas. Cu cât mai elastic, cu atât alergarea e mai eficientă și se economisește energie.').font.size = Pt(11)

p13 = doc.add_paragraph()
p13.add_run('Săriturile antrenează ELASTICITATEA - exact ce necesită alergătorii. Stretching-ul static antrenează FLEXIBILITATEA, necesară dar insuficientă. Combinația celor două este optimă.').font.size = Pt(11)
p13.runs[0].font.bold = True

# PROTOCOLUL COMPLET
doc.add_heading('PROTOCOLUL COMPLET DE RECUPERARE (90 ZILE)', 1)

p14 = doc.add_paragraph()
p14.add_run('Săptămâna 1-8: RESET Fundamental').font.size = Pt(12)
p14.runs[0].font.bold = True

p15 = doc.add_paragraph()
p15.add_run('DIMINEAȚA (20 minute):').font.bold = True
p15.runs[0].font.size = Pt(11)

dim_list = [
    'Decompresie lombară (exercițiu cu minge/roller sub lombar) - 5 min',
    'SĂRITURI VERTICALE - 10 min (1200-1800 sărituri)',
    '  → Suprafață moale (iarbă, covor gros, bandă oprită)',
    '  → Înălțime 5-10cm MAX',
    '  → Landing pe VÂRFURI (forefoot mai întâi)',
    '  → Cadență 120-180 sărituri/minut',
    'Stretching psoas (Thomas Test) - 5 min'
]

for item in dim_list:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()

p16 = doc.add_paragraph()
p16.add_run('DUPĂ-AMIAZĂ (15 minute):').font.bold = True
p16.runs[0].font.size = Pt(11)

dup_list = [
    'Alergare bandă sau iarbă - 10 min MAX',
    '  → Pace 6:00-7:00/km (FĂRĂ viteză!)',
    '  → PAȘI MICI (stride 40-55cm)',
    '  → GENUNCHI JOS (knee lift 20-30cm)',
    '  → Cadență 180-190 pași/min',
    'Sărituri post-alergare - 5 min (opțional)'
]

for item in dup_list:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()

p17 = doc.add_paragraph()
p17.add_run('SEARA (30 minute):').font.bold = True
p17.runs[0].font.size = Pt(11)

seara_list = [
    'Sărituri verticale - 5 min (ultimul ciclu pompare)',
    'Forward fold ASYMMETRIC - 10 min (focus partea mai contractată)',
    'Stretching complet (Figure-4, Cat-Cow, Child\'s pose) - 10 min',
    'Glute activation (bridges, clamshells) - 5 min'
]

for item in seara_list:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.runs[0].font.size = Pt(10)

# MONITORIZARE
doc.add_heading('MONITORIZARE PROGRES', 1)

p18 = doc.add_paragraph()
p18.add_run('Test săptămânal ("3-Minute Reflex Test"):').font.size = Pt(11)
p18.runs[0].font.bold = True

p19 = doc.add_paragraph()
p19.add_run('În fiecare duminică, alergare pe bandă cu tehnica optimă și cronometrare: la ce minut testiculele încep să se strângă? Progresia așteptată:').font.size = Pt(11)

progres = [
    'Săptămâna 0: 3 minute',
    'Săptămâna 2: 5 minute',
    'Săptămâna 4: 8 minute',
    'Săptămâna 6: 12 minute',
    'Săptămâna 8: 15-20 minute sau DELOC'
]

for item in progres:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.runs[0].font.size = Pt(11)

# REZULTATE
doc.add_heading('REZULTATE OBȚINUTE', 1)

p20 = doc.add_paragraph()
p20.add_run('După 8 săptămâni de protocol strict:').font.size = Pt(11)
p20.runs[0].font.bold = True

rezultate = [
    'Testiculele rămân complet relaxate la alergare (15-20 minute fără simptome)',
    'Erecție matinală consistentă 6-7 dimineți/săptămână (vs. 2-3 inițial)',
    'Erecție în context sexual 85-95% din "nivelul mașină" (fără vibrații externe)',
    'Zero disconfort lombar post-alergare',
    'Mobilitate lombară crescută cu ~30%',
    'Forward fold necesar doar 10-15 minute post-alergare (vs. 40 minute inițial)',
    'Capacitate de a alerga 20-30 minute fără simptome (cu tehnica optimă)'
]

for item in rezultate:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.runs[0].font.size = Pt(11)

# CONCLUZII
doc.add_heading('CONCLUZII ȘI RECOMANDĂRI', 1)

p21 = doc.add_paragraph()
p21.add_run('Săriturile verticale ușoare (5-10cm înălțime, pe suprafață moale, aterizare pe vârfuri) sunt superioare alergării pentru sănătatea discurilor lombare din mai multe motive fundamentale:').font.size = Pt(11)

concluzii = [
    'Pompare activă a lichidului în discuri (vs. compresie constantă la alergare)',
    'Decompresie axială pură (vs. shear force la alergare)',
    'Zero solicitare psoas (vs. contractare intensă repetitivă)',
    'Impact controlat, simetric (vs. impact asimetric, variabil)',
    'Antrenament elastic recoil (benefic pentru eficiența alergării ulterioare)'
]

for item in concluzii:
    p = doc.add_paragraph(item, style='List Number')
    p.paragraph_format.left_indent = Inches(0.5)
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()

p22 = doc.add_paragraph()
p22.add_run('IMPORTANT: ').font.bold = True
p22.add_run('Săriturile nu înlocuiesc alergarea pentru beneficiile cardiovasculare și stimularea testosteronului, ci o COMPLETEAZĂ optim. Combinația corectă este:').font.size = Pt(11)

p23 = doc.add_paragraph()
p23.add_run('• 10-20 minute sărituri zilnic (2-3 sesiuni) = pompare discală + elastic recoil').font.size = Pt(11)
doc.add_paragraph()
p24 = doc.add_paragraph()
p24.add_run('• 10-20 minute alergare cu tehnica optimă (pași mici, genunchi jos, pace moderat) = cardio + testosteron').font.size = Pt(11)
doc.add_paragraph()
p25 = doc.add_paragraph()
p25.add_run('• 20-30 minute stretching zilnic (focus psoas) = prevenție contractură').font.size = Pt(11)

doc.add_paragraph()

p26 = doc.add_paragraph()
p26.add_run('Pentru alergători cu probleme lombare, hernie discală, sau disfuncție erectilă asociată cu iritare nervoasă lombară, săriturile verticale reprezintă o intervenție terapeutică simplă, fără costuri, și extrem de eficientă - cu condiția respectării parametrilor corecți (suprafață moale, aterizare pe vârfuri, înălțime mică, frecvență mare).').font.size = Pt(11)
p26.runs[0].font.bold = True

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.add_run('Notă: Acest articol este bazat pe experiență personală și cercetare independentă. Pentru probleme medicale severe, consultați un specialist (fizioterapeut, ortoped, urolog).').font.size = Pt(9)
footer.runs[0].font.italic = True
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Salvare
doc.save('g:/Sarituri_Verticale_Solutie_Lombara_Erectie.docx')
print("Document creat cu succes!")